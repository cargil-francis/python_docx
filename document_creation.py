#imports
from docx import Document
from docx.shared import pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
#create new document
doc = Document()
#add title
title = doc.add_heading ("my new documet sample1", level=2)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
# Add a paragraph with bold and italic text
paragraph=doc.add_paragraph("At times, when you commit some changes to say a remote private repo,The git asks for your credentials so that it can verify if it’s actually you who is committing these changes! This can be simplified by using GIT Credentials. It basically stores your credentials in a safe manner(encrypted) so that you don’t always have to enter your details. So in this tutorial, we will look into how to make workaround easier with GIT Credentials.")
runs=paragraph.run[0]
runs.bold=True
runs.italic=True
# Add a heading
doc.add_heading('Section 1: Introduction', level=2)







